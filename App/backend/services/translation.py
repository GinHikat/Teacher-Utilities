import os
from pathlib import Path
from docx import Document
from deep_translator import GoogleTranslator
try:
    from pdf2docx import Converter
    HAS_PDF2DOCX = True
except ImportError:
    HAS_PDF2DOCX = False
import time

async def convert_pdf_to_docx(pdf_path: Path, progress_callback=None) -> Path:
    """Convert PDF to .docx using pdf2docx."""
    if not HAS_PDF2DOCX:
        raise ImportError("pdf2docx library is not installed for PDF translation.")
    
    if progress_callback:
        await progress_callback(0, 4, "[1/4] Opening document...")
    
    docx_path = pdf_path.with_suffix(".docx")
    cv = Converter(str(pdf_path))
    
    if progress_callback:
        await progress_callback(1, 4, "[2/4] Analyzing document...")
        await progress_callback(2, 4, "[3/4] Parsing pages...")
    
    # Run heavy CPU-bound task in a separate thread to avoid blocking the event loop
    import anyio
    await anyio.to_thread.run_sync(cv.convert, str(docx_path))
    
    if progress_callback:
        await progress_callback(3, 4, "[4/4] Creating pages...")
        await progress_callback(4, 4, "✓ PDF conversion complete.")
        
    cv.close()
    return docx_path

def translate_text(text, target_lang="en"):
    """Translate text to target language."""
    if not text or not text.strip():
        return text or ""

    try:
        # print(f"Translating text to: {target_lang}")
        translated = GoogleTranslator(source="auto", target=target_lang).translate(text)
        return translated if translated is not None else text
    except Exception as e:
        print(f"Translation error: {text[:30]}... -> {e}")
        return text



def translate_paragraph(paragraph, target_lang="en"):
    """Translate a single paragraph while preserving formatting."""
    text = paragraph.text.strip()
    if not text:
        return

    translated = translate_text(text, target_lang=target_lang)
    if translated is None:
        translated = text

    if not paragraph.runs:
        paragraph.add_run(translated)
        return

    # Clear existing runs but keep structure
    for run in paragraph.runs:
        run.text = ""

    paragraph.runs[0].text = translated

async def translate_full_document_async(input_path: Path, output_path: Path, target_lang="en", progress_callback=None):
    if progress_callback:
        await progress_callback(0, 1, f"Initializing translation to: {target_lang}")
    
    if not input_path.exists():
        raise FileNotFoundError(f"File not found: {input_path}")

    # For .pdf files
    is_pdf = False
    temp_docx = None
    if input_path.suffix.lower() == ".pdf":
        is_pdf = True
        temp_docx = await convert_pdf_to_docx(input_path, progress_callback=progress_callback)
        input_path = temp_docx

    try:
        doc = Document(str(input_path))
    except Exception as e:
        raise Exception(f"Cannot open document: {e}")

    total_tasks = len(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_tasks += len(cell.paragraphs)

    if progress_callback:
        await progress_callback(0, total_tasks, f"Analyzing document: {total_tasks} items found.")

    completed = 0
    
    # Translate body paragraphs
    import anyio
    for paragraph in doc.paragraphs:
        await anyio.to_thread.run_sync(translate_paragraph, paragraph, target_lang)
        completed += 1
        if progress_callback:
            await progress_callback(completed, total_tasks, f"Processing body paragraph {completed}...")

    # Translate tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    await anyio.to_thread.run_sync(translate_paragraph, paragraph, target_lang)
                    completed += 1
                    if progress_callback:
                        await progress_callback(completed, total_tasks, f"Processing table cell item {completed}...")

    doc.save(str(output_path))
    
    # Cleanup temporary docx
    if is_pdf and temp_docx and temp_docx.exists():
        try:
            temp_docx.unlink()
        except:
            pass

    return output_path
