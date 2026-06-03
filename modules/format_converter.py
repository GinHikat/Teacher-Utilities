import os
import re
import shutil
from pathlib import Path
try:
    import pythoncom
    import win32com.client
    HAS_WORD = True
except ImportError:
    HAS_WORD = False
from docx import Document
import mammoth
from pdf2docx import Converter

def convert_doc_to_docx(input_path: Path, output_path: Path) -> Path:
    """Convert a .doc file to .docx using Microsoft Word."""
    if not HAS_WORD:
        raise RuntimeError("Microsoft Word automation (.doc to .docx) is only supported on Windows hosts.")
    pythoncom.CoInitialize()
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(input_path.resolve()))
        # FileFormat=16 is for wdFormatXMLDocument (.docx)
        doc.SaveAs(str(output_path.resolve()), FileFormat=16)
        doc.Close()
        word.Quit()
    except Exception as e:
        raise RuntimeError(f"Error converting .doc to .docx: {e}")
    finally:
        pythoncom.CoUninitialize()
    return output_path

def convert_docx_to_doc(input_path: Path, output_path: Path) -> Path:
    """Convert a .docx file to .doc using Microsoft Word."""
    if not HAS_WORD:
        raise RuntimeError("Microsoft Word automation (.docx to .doc) is only supported on Windows hosts.")
    pythoncom.CoInitialize()
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(input_path.resolve()))
        # FileFormat=0 is for wdFormatDocument (.doc)
        doc.SaveAs(str(output_path.resolve()), FileFormat=0)
        doc.Close()
        word.Quit()
    except Exception as e:
        raise RuntimeError(f"Error converting .docx to .doc: {e}")
    finally:
        pythoncom.CoUninitialize()
    return output_path

def convert_docx_to_pdf(input_path: Path, output_path: Path) -> Path:
    """Convert a .docx file to .pdf using Microsoft Word."""
    if not HAS_WORD:
        raise RuntimeError("Microsoft Word automation (.docx to .pdf) is only supported on Windows hosts.")
    pythoncom.CoInitialize()
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(input_path.resolve()))
        # FileFormat=17 is for wdFormatPDF (.pdf)
        doc.SaveAs(str(output_path.resolve()), FileFormat=17)
        doc.Close()
        word.Quit()
    except Exception as e:
        raise RuntimeError(f"Error converting .docx to .pdf: {e}")
    finally:
        pythoncom.CoUninitialize()
    return output_path

def convert_pdf_to_docx(input_path: Path, output_path: Path) -> Path:
    """Convert a .pdf file to .docx using pdf2docx."""
    try:
        cv = Converter(str(input_path.resolve()))
        cv.convert(str(output_path.resolve()), start=0, end=None)
        cv.close()
    except Exception as e:
        raise RuntimeError(f"Error converting .pdf to .docx: {e}")
    return output_path

def convert_docx_to_markdown(input_path: Path, output_path: Path) -> Path:
    """Convert a .docx file to Markdown using mammoth."""
    try:
        with open(input_path, "rb") as f:
            result = mammoth.convert_to_markdown(f)
            markdown_text = result.value
            
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(markdown_text)
    except Exception as e:
        raise RuntimeError(f"Error converting .docx to Markdown: {e}")
    return output_path

def parse_inline_formatting(paragraph, text: str):
    """Parse basic inline markdown (bold, italic, inline code) and add runs to the paragraph."""
    # Pattern to match bold (**, __), italic (*, _), and inline code (`)
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|__.*?__|___.*?___|_.*?_|`.*?`)')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
        elif part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('___') and part.endswith('___'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('__') and part.endswith('__'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('_') and part.endswith('_'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

def convert_markdown_to_docx(input_path: Path, output_path: Path) -> Path:
    """Convert a Markdown file to .docx."""
    try:
        with open(input_path, "r", encoding="utf-8") as f:
            content = f.read()
            
        doc = Document()
        lines = content.splitlines()
        
        in_list = False
        
        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                in_list = False
                continue
                
            # Headings
            if stripped.startswith('#'):
                m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
                if m:
                    level = len(m.group(1))
                    text = m.group(2)
                    p = doc.add_heading('', level=level)
                    parse_inline_formatting(p, text)
                    in_list = False
                    continue
            
            # Numbered lists
            if re.match(r'^\d+\.\s+', stripped):
                text = re.sub(r'^\d+\.\s+', '', stripped)
                p = doc.add_paragraph('', style='List Number')
                parse_inline_formatting(p, text)
                in_list = True
                continue
                
            # Bullet lists
            if stripped.startswith(('- ', '* ', '+ ')):
                text = stripped[2:]
                p = doc.add_paragraph('', style='List Bullet')
                parse_inline_formatting(p, text)
                in_list = True
                continue
                
            # Blockquotes
            if stripped.startswith('> '):
                text = stripped[2:]
                p = doc.add_paragraph('', style='Quote')
                parse_inline_formatting(p, text)
                in_list = False
                continue
                
            # Normal paragraph
            p = doc.add_paragraph('')
            parse_inline_formatting(p, stripped)
            in_list = False
            
        doc.save(str(output_path.resolve()))
    except Exception as e:
        raise RuntimeError(f"Error converting Markdown to .docx: {e}")
    return output_path

def convert_file(input_path: Path, output_path: Path) -> Path:
    """
    Main orchestrator for file conversion. Automatically resolves
    multi-step conversions through an intermediate .docx state.
    """
    src_ext = input_path.suffix.lower()
    dest_ext = output_path.suffix.lower()
    
    if src_ext == dest_ext:
        shutil.copy(str(input_path), str(output_path))
        return output_path
        
    # Standard single conversions:
    if src_ext == '.doc' and dest_ext == '.docx':
        return convert_doc_to_docx(input_path, output_path)
        
    if src_ext == '.docx' and dest_ext == '.doc':
        return convert_docx_to_doc(input_path, output_path)
        
    if src_ext == '.docx' and dest_ext == '.pdf':
        return convert_docx_to_pdf(input_path, output_path)
        
    if src_ext == '.pdf' and dest_ext == '.docx':
        return convert_pdf_to_docx(input_path, output_path)
        
    if src_ext == '.docx' and dest_ext == '.md':
        return convert_docx_to_markdown(input_path, output_path)
        
    if src_ext == '.md' and dest_ext == '.docx':
        return convert_markdown_to_docx(input_path, output_path)
        
    # Multi-step conversions (utilizing a temporary .docx intermediate file):
    # 1. Source is .doc (which Word can natively handle, but we standardize as .docx)
    if src_ext == '.doc':
        temp_docx = input_path.with_suffix('.temp_doc.docx')
        convert_doc_to_docx(input_path, temp_docx)
        try:
            return convert_file(temp_docx, output_path)
        finally:
            if temp_docx.exists():
                temp_docx.unlink()
                
    # 2. Source is .pdf converting to doc/pdf/md
    if src_ext == '.pdf':
        temp_docx = input_path.with_suffix('.temp_pdf.docx')
        convert_pdf_to_docx(input_path, temp_docx)
        try:
            return convert_file(temp_docx, output_path)
        finally:
            if temp_docx.exists():
                temp_docx.unlink()
                
    # 3. Source is .md converting to doc/pdf
    if src_ext == '.md':
        temp_docx = input_path.with_suffix('.temp_md.docx')
        convert_markdown_to_docx(input_path, temp_docx)
        try:
            return convert_file(temp_docx, output_path)
        finally:
            if temp_docx.exists():
                temp_docx.unlink()
                
    # 4. Destination is .doc (already handled docx -> doc above; others convert to docx first)
    if dest_ext == '.doc':
        temp_docx = output_path.with_suffix('.temp_out.docx')
        convert_file(input_path, temp_docx)
        try:
            return convert_docx_to_doc(temp_docx, output_path)
        finally:
            if temp_docx.exists():
                temp_docx.unlink()
                
    # 5. Destination is .pdf (already handled docx -> pdf above; others convert to docx first)
    if dest_ext == '.pdf':
        temp_docx = output_path.with_suffix('.temp_out.docx')
        convert_file(input_path, temp_docx)
        try:
            return convert_docx_to_pdf(temp_docx, output_path)
        finally:
            if temp_docx.exists():
                temp_docx.unlink()
                
    # 6. Destination is .md (already handled docx -> md above; others convert to docx first)
    if dest_ext == '.md':
        temp_docx = output_path.with_suffix('.temp_out.docx')
        convert_file(input_path, temp_docx)
        try:
            return convert_docx_to_markdown(temp_docx, output_path)
        finally:
            if temp_docx.exists():
                temp_docx.unlink()

    raise ValueError(f"Unsupported conversion from {src_ext} to {dest_ext}")
